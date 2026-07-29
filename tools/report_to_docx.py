"""Render report/REPORT.md to a .docx for submission.

Pandoc is not installed on the build machine, and this document has three properties a naive
dump would destroy: markdown pipe tables (the report leans on them for every measurement),
mixed Arabic/Latin paragraphs that need per-paragraph direction, and inline `code` spans that
must stay distinguishable from prose.

Usage:  python tools/report_to_docx.py [-o report/REPORT.docx]
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "REPORT.md"

_ARABIC = re.compile(r"[؀-ۿ]")
_LATIN = re.compile(r"[A-Za-z]")
# **bold**, `code`, [text](url) — captured in one pass so a bold span containing code still splits
_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def is_rtl(text: str) -> bool:
    """Direction is decided per paragraph, the same script-ratio rule the agent uses (FR1)."""
    return len(_ARABIC.findall(text)) > len(_LATIN.findall(text))


def add_runs(paragraph, text: str) -> None:
    """Write inline markdown as formatted runs rather than literal asterisks and backticks."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("[") and "](" in part:
            label, _, url = part[1:].partition("](")
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x18, 0x4A, 0x90)
            run.underline = True
            # the URL is kept visible: a printed report loses clickable links entirely
            paragraph.add_run(f" ({url.rstrip(')')})").font.size = Pt(8)
        else:
            paragraph.add_run(part)


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)


def emit_paragraph(doc, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    add_runs(p, text)
    if is_rtl(text):
        set_rtl(p)


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def emit_table(doc, rows: list[str]) -> None:
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]          # rows[1] is the --- separator
    width = max([len(header)] + [len(r) for r in body])

    table = doc.add_table(rows=1, cols=width)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = ""
        add_runs(cell.paragraphs[0], header[i] if i < len(header) else "")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            cell.text = ""
            value = row[i] if i < len(row) else ""
            add_runs(cell.paragraphs[0], value)
            if is_rtl(value):
                set_rtl(cell.paragraphs[0])


_BLOCK_START = re.compile(r"^\s*(#{1,6}\s|>\s|[-*+]\s|\d+[.)]\s|\||```|(-{3,}|_{3,}|\*{3,})\s*$)")


def unwrap(lines: list[str]) -> list[str]:
    """Join hard-wrapped source lines back into whole paragraphs.

    REPORT.md is wrapped at ~100 characters. Rendering it line by line produced a Word paragraph
    break every ~100 characters mid-sentence, and split bold spans across lines so «**not on
    Dawlati\\nat all**» reached the page as literal asterisks — 20 such leaks before this existed.
    Continuation lines are folded into the block above; anything that starts a new block (heading,
    list item, quote, table row, fence, rule) or a blank line ends it.
    """
    out: list[str] = []
    buffer: list[str] = []
    in_code = False

    def flush() -> None:
        if buffer:
            out.append(" ".join(x.strip() for x in buffer))
            buffer.clear()

    for line in lines:
        if line.strip().startswith("```"):
            flush()
            out.append(line)
            in_code = not in_code
            continue
        if in_code:
            out.append(line)
            continue
        if not line.strip():
            flush()
            out.append("")
            continue
        if _BLOCK_START.match(line):
            # a wrapped LIST item continues its bullet; every other block type starts fresh
            if buffer and re.match(r"^\s*([-*+]|\d+[.)])\s", out[-1] if out else ""):
                pass
            flush()
            buffer.append(line)
            # tables and rules are single-line blocks — emit immediately
            if re.match(r"^\s*(\||```|(-{3,}|_{3,}|\*{3,})\s*$)", line):
                flush()
            continue
        buffer.append(line)
    flush()
    return out


def convert(source: Path, target: Path) -> int:
    lines = unwrap(source.read_text(encoding="utf-8").splitlines())
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i, n_tables, in_code = 0, 0, False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(18)
                code_buffer, in_code = [], False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        # pipe table: a header row followed by a |---|---| separator
        if (stripped.startswith("|") and i + 1 < len(lines)
                and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1])):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            emit_table(doc, block)
            doc.add_paragraph()
            n_tables += 1
            continue

        if not stripped:
            i += 1
            continue
        if re.match(r"^(-{3,}|_{3,}|\*{3,})$", stripped):
            # A markdown rule becomes whitespace, not a drawn line. Rendering it as a row of dash
            # characters put a decorative divider between every section, which is a visual tic the
            # headings already handle.
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            emit_paragraph(doc, stripped.lstrip("# ").strip(), style=f"Heading {min(level, 4)}")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, stripped[2:])
        elif re.match(r"^[-*+]\s+", stripped):
            emit_paragraph(doc, re.sub(r"^[-*+]\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            emit_paragraph(doc, re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
        else:
            emit_paragraph(doc, stripped)
        i += 1

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    print(f"wrote {target}  ({len(lines)} md lines, {n_tables} tables)")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("-o", "--out", default=str(ROOT / "report" / "REPORT.docx"))
    ap.add_argument("-i", "--input", default=str(SOURCE))
    args = ap.parse_args()
    return convert(Path(args.input), Path(args.out))


if __name__ == "__main__":
    raise SystemExit(main())
